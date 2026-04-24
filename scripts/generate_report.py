"""
Word 레포트 생성 스크립트.
report/template/ 의 .docx 템플릿에서 {{PLACEHOLDER}} 를 치환하여
report/output/final_report.docx 를 생성한다.

사용법:
    python scripts/generate_report.py

의존성:
    pip install python-docx
"""

import glob
import os
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

BASE_DIR = Path(__file__).parent.parent
TEMPLATE_DIR = BASE_DIR / "report" / "template"
OUTPUT_DIR = BASE_DIR / "report" / "output"
OUTPUT_FILE = OUTPUT_DIR / "final_report.docx"

# ── 템플릿 주제가 확정되면 여기에 값을 채운다 ──────────────────────────────
PLACEHOLDERS: dict[str, str] = {
    # "{{TITLE}}": "과제 제목",
    # "{{AUTHOR}}": "김수진",
    # "{{DATE}}": "2026-04-24",
    # "{{SECTION_BACKGROUND}}": "...",
}
# ──────────────────────────────────────────────────────────────────────────────


def _replace_in_paragraph(para, mapping: dict[str, str]) -> None:
    """단락(paragraph) 안의 run들을 합쳐서 플레이스홀더를 치환한다."""
    full_text = "".join(run.text for run in para.runs)
    replaced = full_text
    for key, value in mapping.items():
        replaced = replaced.replace(key, value)
    if replaced != full_text:
        for i, run in enumerate(para.runs):
            run.text = replaced if i == 0 else ""


def _replace_in_table(table, mapping: dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _replace_in_paragraph(para, mapping)


def generate(template_path: Path, output_path: Path, mapping: dict[str, str]) -> None:
    doc = Document(str(template_path))

    for para in doc.paragraphs:
        _replace_in_paragraph(para, mapping)

    for table in doc.tables:
        _replace_in_table(table, mapping)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"[OK] 레포트 생성 완료: {output_path}")


def find_template() -> Path:
    candidates = sorted(TEMPLATE_DIR.glob("*.docx"))
    if not candidates:
        raise FileNotFoundError(f"템플릿 파일이 없습니다: {TEMPLATE_DIR}")
    if len(candidates) > 1:
        print(f"[WARN] 템플릿 파일이 여러 개입니다. 첫 번째 파일을 사용합니다: {candidates[0].name}")
    return candidates[0]


if __name__ == "__main__":
    if not PLACEHOLDERS:
        print("[WARN] PLACEHOLDERS 딕셔너리가 비어 있습니다. 템플릿 주제 확정 후 채워주세요.")

    template = find_template()
    generate(template, OUTPUT_FILE, PLACEHOLDERS)
